'''
    Quick play to explore the docx module and how docx docs are encoded
    Using CoupaMarkUp style to indicate some processing instruction

'''

import docx
import copy



def printProcessingInstructions(doc):
    """
    Print out all the template processing instructions.  These are all in CoupaMarkUp style in the document
    :param doc: the xml doc object
    :return: None
    """
    for p in doc.paragraphs:
        if p.style.name == 'CoupaMarkUp':
            print(f'Processing Instruction Paragraph: {p.text}')
        inst = ''
        for i, r in enumerate(p.runs):
            # merge adjacent mark up runs together
            if r.style.name.find('CoupaMarkUp') != -1:
                inst = inst + r.text
            else:
                if inst != '':  # end of adjacent instructions
                    print(f"Processing Instruction run: {inst}")
                    inst = ''
        if inst != '':  # inst may have been the last run in the paragraph
            print(f"Processing Instruction run: {inst}")

def replaceVariableRunInstruction(paragraph, instruction_start_run, instruction_end_run, value):
    # add the variable text to end of the previous run and empty the other instruction runs
    # or if this is the first run add it to the start of the next none instruction run

    # maintain and the white space in the process instruction in the replacement string
    full_instruction_text = ''
    for i in range(instruction_start_run, instruction_end_run + 1):
        full_instruction_text = full_instruction_text + paragraph.runs[i].text

    leading_white_space = ''
    trailing_white_space = ''

    len_leading_space = len(full_instruction_text) - len(full_instruction_text.lstrip())
    if len_leading_space > 0:
        leading_white_space = full_instruction_text[:len_leading_space]

    len_trailing_space = len(full_instruction_text) - len(full_instruction_text.rstrip())
    if len_trailing_space > 0:
        trailing_white_space = full_instruction_text[-len_trailing_space:]

    replacement = leading_white_space + value + trailing_white_space

    if instruction_start_run != 0:
        paragraph.runs[instruction_start_run - 1].text = paragraph.runs[instruction_start_run - 1].text + replacement
    else:
        paragraph.runs[instruction_end_run + 1].text = replacement + paragraph.runs[instruction_end_run + 1].text

    # TODO what if the instruction runs are the only runs

    for j in range(instruction_start_run, instruction_end_run + 1):  # remember range doesn't include last value so +1
        paragraph.runs[j].text = ''


def processVariables(doc, variables):
    for p in doc.paragraphs:
        if p.style.name == 'CoupaMarkUp':
            print(f'Processing Instruction Paragraph: {p.text}')

            # simple dictionary look up for now
            if (p.text.strip() in variables):
                p.text = p.text.replace(p.text.strip(), variables[p.text])

        inst = ''
        instStartRun = None
        instEndRun = None
        for i, r in enumerate(p.runs):
            # merge adjacent mark up runs together to form one instruction
            if r.style.name.find('CoupaMarkUp') != -1:
                if instStartRun == None:
                    instStartRun = i
                inst = inst + r.text
            else:
                if instStartRun != None:  # end of adjacent instructions
                    instEndRun = i-1
                    print(f"Processing Instruction run: {inst}")

                    # simple dictionary look up for now
                    if (inst.strip() in variables):
                        replaceVariableRunInstruction(p, instStartRun, instEndRun, variables[inst.strip()])

                    inst = ''
                    instStartRun = None
                    instEndRun = None
        if instStartRun != None:  # inst may have been the last run in the paragraph
            instEndRun = i
            print(f"Processing Instruction run: {inst}")

            # simple dictionary look up for now
            if (inst.strip() in variables):
                replaceVariableRunInstruction(p, instStartRun, instEndRun, variables[inst.strip()])

            print(f"Processing Instruction run: {inst}")


def processTables(doc):
    """
    Got through all the cells in all the tables in the doc and process
    any CoupaMarkup in them
    :param doc: the docx doc object to process
    :return: None
    """
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if p.style.name == 'CoupaMarkUp':
                        print(f'Processing Instruction Paragraph: {p.text}')
                        p.text = "PROCESSED"


def processRepeatParagraphs(doc):
    processing_instruction = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'CoupaMarkUp':
            print(f'Processing Instruction Paragraph: {p.text}')
            processing_instruction = p.text.strip()

            if (processing_instruction.find('repeat_') == 0):
                repeat_start_paragraph = i
                # rest of the instruction is the number of times to repeat
                times_str = processing_instruction[7:len(processing_instruction)]
                times = int(times_str)

                # find the end of the instruction
                for j in range(repeat_start_paragraph + 1, len(doc.paragraphs)):
                    par = doc.paragraphs[j]
                    if par.style.name == 'CoupaMarkUp':
                        print(f'Processing Instruction Paragraph: {p.text}')
                        processing_instruction = par.text.strip()

                        if (processing_instruction.find('end') == 0):
                            repeat_end_paragraph = j-1

                # Copy the paragraphs in between...

                # Just try one for now
                copy_from_paragrph = doc.paragraphs[repeat_start_paragraph + 1]
                end_instruction_paragraph = doc.paragraphs[repeat_end_paragraph + 1]
                end_instruction_paragraph.insert_paragraph_before("Inserted paragraph")

                # the new paragraph will have been inserted where the old end was
                doc.paragraphs[repeat_end_paragraph + 1] = copy_from_paragrph

                print(f'Copy paragraph from {repeat_start_paragraph} to {repeat_end_paragraph} {times} times.')


doc = docx.Document('test.docx')
# printProcessingInstructions(doc)

vardict = {"Val_1" : "Replacement Value",
           "clause_liability" : "The rain is liable for this",
           "change_me" : "I've been changed",
           "change_me_2" : "String with line break \nNext Line"}
processVariables(doc, vardict)
processTables(doc)
processRepeatParagraphs(doc)

doc.save('test_out.docx')